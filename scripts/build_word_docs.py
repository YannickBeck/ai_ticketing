from __future__ import annotations

import re
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORD_DIR = ROOT / "docs" / "word"

SOURCES = [
    (
        WORD_DIR / "spargelstand-executive-konzept.md",
        WORD_DIR / "Spargelstand-App-Executive-Konzept.docx",
        "Executive-Konzept",
    ),
    (
        WORD_DIR / "spargelstand-technisches-konzept.md",
        WORD_DIR / "Spargelstand-App-Technisches-Konzept.docx",
        "Technisches Konzept",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    if bold:
        run.font.color.rgb = RGBColor(255, 255, 255)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`[^`]*`|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if link:
                paragraph.add_run(f"{link.group(1)} ({link.group(2)})")
            else:
                paragraph.add_run(part)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 24, RGBColor(31, 78, 121)),
        ("Heading 1", 17, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11, RGBColor(31, 78, 121)),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    styles["List Bullet"].font.name = "Aptos"
    styles["List Number"].font.name = "Aptos"


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Spargelstand-App")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Stand: 22. Mai 2026").italic = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Konzeptdokument zur MVP-Umsetzung eines standortbasierten Preorder- und Pickup-Marktplatzes "
        "für dezentrale landwirtschaftliche Verkaufsstände."
    )
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|") and lines[idx].strip().endswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1

    rows = []
    for raw in table_lines:
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            is_header = row_idx == 0
            set_cell_text(cell, text, bold=is_header)
            if is_header:
                set_cell_shading(cell, "1F4E79")
            else:
                set_cell_shading(cell, "F7F9FB" if row_idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph()


def add_code_block(doc: Document, code: str, language: str | None) -> None:
    if language == "mermaid":
        paragraph = doc.add_paragraph()
        paragraph.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Normal"
        paragraph.add_run(
            "Diagrammhinweis: Das Mermaid-Diagramm ist in der Markdown-Quelle enthalten. "
            "Im Word-Dokument wird der Inhalt durch die angrenzenden Prozesslisten und Tabellen beschrieben."
        ).italic = True
        return

    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 0
    paragraph_buffer: list[str] = []
    in_code = False
    code_language: str | None = None
    code_lines: list[str] = []
    skipped_first_h1 = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline_runs(paragraph, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                flush_paragraph()
                in_code = True
                code_language = stripped[3:].strip() or None
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines), code_language)
                code_language = None
                code_lines = []
            idx += 1
            continue

        if in_code:
            code_lines.append(raw)
            idx += 1
            continue

        if not stripped:
            flush_paragraph()
            idx += 1
            continue

        if stripped == "---":
            flush_paragraph()
            doc.add_paragraph()
            idx += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
            else:
                doc.add_heading(text, level=min(level, 3))
            idx += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet.group(1))
            idx += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, ordered.group(1))
            idx += 1
            continue

        paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph()


def build_docx(source: Path, target: Path, subtitle: str) -> None:
    markdown = source.read_text(encoding="utf-8")
    title_match = re.search(r"^#\s+(.+)$", markdown, flags=re.MULTILINE)
    title = title_match.group(1).strip() if title_match else source.stem

    doc = Document()
    configure_styles(doc)
    add_cover(doc, title, subtitle)
    render_markdown(doc, markdown)

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Spargelstand-App - MVP-Konzept")

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)

    if not zipfile.is_zipfile(target):
        raise RuntimeError(f"DOCX wurde nicht als gültiges Office-Open-XML-Archiv erzeugt: {target}")


def main() -> None:
    for source, target, subtitle in SOURCES:
        build_docx(source, target, subtitle)
        print(f"created {target.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
